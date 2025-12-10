#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de Templates Word (.docx) para Documentos Técnicos

Este script cria templates profissionais formatados para:
- ETP (Estudo Técnico Preliminar)
- TR (Termo de Referência)
- MD (Memorial Descritivo)

Com placeholders prontos para substituição automática.

Autor: Sistema de Automação - Prefeitura Municipal
Data: Dezembro 2025
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def configurar_estilos(doc):
    """Configura estilos padrão do documento."""
    # Estilo para título principal
    try:
        titulo_style = doc.styles['Heading 1']
    except:
        titulo_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    titulo_style.font.name = 'Arial'
    titulo_style.font.size = Pt(16)
    titulo_style.font.bold = True
    titulo_style.font.color.rgb = RGBColor(0, 0, 139)  # Azul escuro
    
    # Estilo para subtítulos
    try:
        subtitulo_style = doc.styles['Heading 2']
    except:
        subtitulo_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    subtitulo_style.font.name = 'Arial'
    subtitulo_style.font.size = Pt(14)
    subtitulo_style.font.bold = True
    subtitulo_style.font.color.rgb = RGBColor(0, 51, 102)

def adicionar_cabecalho(doc, titulo):
    """Adiciona cabeçalho oficial do documento."""
    # Logo/Brasão (placeholder - você pode adicionar imagem depois)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PREFEITURA MUNICIPAL DE NOVA PETRÓPOLIS\n")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    
    run = p.add_run("ESTADO DO RIO GRANDE DO SUL\n")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    run = p.add_run("Secretaria de Obras e Infraestrutura")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espaço
    
    # Título do documento
    p = doc.add_heading(titulo, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espaço

def criar_template_etp(caminho_saida):
    """Cria template para Estudo Técnico Preliminar."""
    doc = Document()
    configurar_estilos(doc)
    
    adicionar_cabecalho(doc, "ESTUDO TÉCNICO PRELIMINAR - ETP")
    
    # Seção 1: Identificação
    doc.add_heading("1. IDENTIFICAÇÃO", level=2)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    dados_identificacao = [
        ("Objeto:", "{{OBJETO}}"),
        ("Município:", "{{LOCAL}}"),
        ("Valor Estimado:", "R$ {{VALOR_GLOBAL}}"),
        ("Valor Repasse:", "R$ {{VALOR_REPASSE}}"),
        ("Contrapartida:", "R$ {{VALOR_CONTRAPARTIDA}}"),
        ("Área Total:", "{{AREA_TOTAL}} m²"),
        ("Data Base:", "{{DATA_BASE}}"),
        ("Responsável Técnico:", "{{RESPONSAVEL}}")
    ]
    
    for i, (campo, valor) in enumerate(dados_identificacao):
        row = table.rows[i]
        row.cells[0].text = campo
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = valor
    
    doc.add_paragraph()
    
    # Seção 2: Justificativa
    doc.add_heading("2. JUSTIFICATIVA E DESCRIÇÃO DA NECESSIDADE", level=2)
    
    p = doc.add_paragraph()
    p.add_run("{{TEXTO_IA}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Justificativa complementar: {{JUSTIFICATIVA}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 3: Estimativa de Custos
    doc.add_heading("3. ESTIMATIVA DE CUSTOS E CRONOGRAMA", level=2)
    
    p = doc.add_paragraph()
    texto_custos = f"""O valor global estimado para a execução do objeto descrito é de R$ {{{{VALOR_GLOBAL}}}}, conforme planilha orçamentária detalhada em anexo, elaborada com base nos custos referenciais do Sistema Nacional de Pesquisa de Custos e Índices da Construção Civil (SINAPI), referência {{{{DATA_BASE}}}}.

O investimento será composto por:
- Repasse Federal: R$ {{{{VALOR_REPASSE}}}}
- Contrapartida Municipal: R$ {{{{VALOR_CONTRAPARTIDA}}}}

BDI aplicado: {{{{BDI}}}}%"""
    
    p.add_run(texto_custos)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 4: Conformidade Legal
    doc.add_heading("4. CONFORMIDADE COM A LEI Nº 14.133/2021", level=2)
    
    p = doc.add_paragraph()
    texto_legal = """Este Estudo Técnico Preliminar foi elaborado em conformidade com o disposto no art. 18, inciso I, da Lei Federal nº 14.133/2021 (Nova Lei de Licitações e Contratos Administrativos), visando demonstrar:

a) A caracterização do interesse público envolvido;
b) A especificação do objeto com todos os elementos que o caracterizam;
c) A justificativa da contratação em face da necessidade pública;
d) As estimativas de quantidades e custos;
e) A compatibilidade com o Plano Plurianual e com a Lei de Diretrizes Orçamentárias."""
    
    p.add_run(texto_legal)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nova Petrópolis/RS, {{{{DATA_ATUAL}}}}\n\n")
    
    p.add_run("_" * 50 + "\n")
    p.add_run("{{RESPONSAVEL}}\n")
    p.add_run("{{SETOR}}")
    
    # Salvar
    doc.save(caminho_saida)
    print(f"✅ Template ETP criado: {caminho_saida}")

def criar_template_tr(caminho_saida):
    """Cria template para Termo de Referência."""
    doc = Document()
    configurar_estilos(doc)
    
    adicionar_cabecalho(doc, "TERMO DE REFERÊNCIA")
    
    # Seção 1: Objeto
    doc.add_heading("1. DO OBJETO", level=2)
    p = doc.add_paragraph()
    p.add_run("Constitui objeto do presente Termo de Referência a contratação de: {{OBJETO}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 2: Justificativa
    doc.add_heading("2. DA JUSTIFICATIVA", level=2)
    p = doc.add_paragraph()
    p.add_run("{{TEXTO_IA}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("{{JUSTIFICATIVA}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 3: Especificações Técnicas
    doc.add_heading("3. DAS ESPECIFICAÇÕES TÉCNICAS E QUANTITATIVOS", level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    specs = [
        ("Local de Execução:", "{{LOCAL}}"),
        ("Área Total a Intervir:", "{{AREA_TOTAL}} m²"),
        ("Data Base do Orçamento:", "{{DATA_BASE}}"),
        ("BDI Aplicado:", "{{BDI}}%")
    ]
    
    for i, (campo, valor) in enumerate(specs):
        row = table.rows[i]
        row.cells[0].text = campo
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = valor
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("As especificações técnicas detalhadas, planilhas de quantitativos e composições de custos encontram-se em anexo a este Termo de Referência.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 4: Obrigações da Contratada
    doc.add_heading("4. DAS OBRIGAÇÕES DA CONTRATADA", level=2)
    
    obrigacoes = [
        "Executar fielmente o objeto contratado, em conformidade com as normas técnicas vigentes;",
        "Manter durante toda a execução do contrato as condições de habilitação e qualificação exigidas;",
        "Responsabilizar-se pelos encargos trabalhistas, previdenciários, fiscais e comerciais;",
        "Reparar, corrigir ou substituir, às suas expensas, no todo ou em parte, o objeto em que se verificarem vícios;",
        "Fornecer Anotação de Responsabilidade Técnica (ART) ou Registro de Responsabilidade Técnica (RRT)."
    ]
    
    for obrigacao in obrigacoes:
        doc.add_paragraph(obrigacao, style='List Bullet')
    
    # Seção 5: Critérios de Medição
    doc.add_heading("5. DOS CRITÉRIOS DE MEDIÇÃO E PAGAMENTO", level=2)
    p = doc.add_paragraph()
    texto_medicao = """Os serviços serão medidos mensalmente conforme os quantitativos efetivamente executados e aceitos pela fiscalização, em conformidade com a planilha orçamentária e o cronograma físico-financeiro.

O pagamento será realizado em até 30 (trinta) dias após a apresentação da medição e respectiva nota fiscal, devidamente atestada pelo fiscal do contrato."""
    p.add_run(texto_medicao)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 6: Valor Estimado
    doc.add_heading("6. DO VALOR ESTIMADO", level=2)
    
    p = doc.add_paragraph()
    texto_valor = f"""O valor total estimado para a contratação é de R$ {{{{VALOR_GLOBAL}}}} (valor por extenso), conforme detalhamento:

• Valor de Repasse Federal: R$ {{{{VALOR_REPASSE}}}}
• Contrapartida Municipal: R$ {{{{VALOR_CONTRAPARTIDA}}}}

Os preços foram estimados com base no Sistema Nacional de Pesquisa de Custos e Índices da Construção Civil (SINAPI), referência {{{{DATA_BASE}}}}, acrescidos de BDI de {{{{BDI}}}}%."""
    
    p.add_run(texto_valor)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 7: Prazo de Execução
    doc.add_heading("7. DO PRAZO DE EXECUÇÃO", level=2)
    p = doc.add_paragraph()
    p.add_run("O prazo para execução dos serviços será de [INSERIR PRAZO] dias corridos, contados a partir da emissão da Ordem de Serviço.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 8: Fundamentação Legal
    doc.add_heading("8. DA FUNDAMENTAÇÃO LEGAL", level=2)
    p = doc.add_paragraph()
    texto_legal = """Este Termo de Referência foi elaborado em conformidade com:
• Lei Federal nº 14.133/2021 (Nova Lei de Licitações e Contratos Administrativos);
• Lei Complementar nº 101/2000 (Lei de Responsabilidade Fiscal);
• Normas técnicas da ABNT aplicáveis."""
    p.add_run(texto_legal)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nova Petrópolis/RS, {{{{DATA_ATUAL}}}}\n\n")
    p.add_run("_" * 50 + "\n")
    p.add_run("{{RESPONSAVEL}}\n")
    p.add_run("{{SETOR}}")
    
    doc.save(caminho_saida)
    print(f"✅ Template TR criado: {caminho_saida}")

def criar_template_md(caminho_saida):
    """Cria template para Memorial Descritivo."""
    doc = Document()
    configurar_estilos(doc)
    
    adicionar_cabecalho(doc, "MEMORIAL DESCRITIVO")
    
    # Identificação
    doc.add_heading("IDENTIFICAÇÃO DO PROJETO", level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    id_projeto = [
        ("Projeto:", "{{OBJETO}}"),
        ("Localização:", "{{LOCAL}}"),
        ("Área de Intervenção:", "{{AREA_TOTAL}} m²"),
        ("Data Base:", "{{DATA_BASE}}"),
        ("Valor Estimado:", "R$ {{VALOR_GLOBAL}}"),
        ("Responsável Técnico:", "{{RESPONSAVEL}}")
    ]
    
    for i, (campo, valor) in enumerate(id_projeto):
        row = table.rows[i]
        row.cells[0].text = campo
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = valor
    
    # Seção 1: Considerações Iniciais
    doc.add_heading("1. CONSIDERAÇÕES INICIAIS", level=2)
    p = doc.add_paragraph()
    p.add_run("{{TEXTO_IA}}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 2: Descrição do Projeto
    doc.add_heading("2. DESCRIÇÃO GERAL DO PROJETO", level=2)
    p = doc.add_paragraph()
    texto_descricao = """O presente Memorial Descritivo tem por finalidade detalhar os aspectos técnicos e construtivos do projeto: {{OBJETO}}.

Justificativa: {{JUSTIFICATIVA}}

A intervenção abrange uma área total de {{AREA_TOTAL}} m², localizada em {{LOCAL}}, com investimento global estimado em R$ {{VALOR_GLOBAL}}, sendo R$ {{VALOR_REPASSE}} provenientes de repasse federal e R$ {{VALOR_CONTRAPARTIDA}} de contrapartida municipal."""
    
    p.add_run(texto_descricao)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 3: Especificações Técnicas
    doc.add_heading("3. ESPECIFICAÇÕES TÉCNICAS DOS SERVIÇOS", level=2)
    
    doc.add_heading("3.1 Serviços Preliminares", level=3)
    p = doc.add_paragraph()
    p.add_run("[Descrever serviços preliminares: canteiro, mobilização, locação, etc.]")
    
    doc.add_heading("3.2 Serviços de Terraplanagem", level=3)
    p = doc.add_paragraph()
    p.add_run("[Descrever serviços de terraplanagem se aplicável]")
    
    doc.add_heading("3.3 Serviços Principais", level=3)
    p = doc.add_paragraph()
    p.add_run("[Detalhar execução dos serviços principais do projeto]")
    
    doc.add_heading("3.4 Serviços Complementares", level=3)
    p = doc.add_paragraph()
    p.add_run("[Descrever serviços complementares e acabamentos]")
    
    # Seção 4: Normas e Padrões
    doc.add_heading("4. NORMAS TÉCNICAS APLICÁVEIS", level=2)
    
    normas = [
        "Normas Técnicas da Associação Brasileira de Normas Técnicas (ABNT);",
        "Especificações do Departamento Nacional de Infraestrutura de Transportes (DNIT);",
        "Diretrizes da Caixa Econômica Federal para projetos de infraestrutura;",
        "Código de Obras Municipal;",
        "Legislação ambiental vigente."
    ]
    
    for norma in normas:
        doc.add_paragraph(norma, style='List Bullet')
    
    # Seção 5: Orçamento
    doc.add_heading("5. ORÇAMENTO E CRONOGRAMA", level=2)
    p = doc.add_paragraph()
    texto_orcamento = f"""O orçamento detalhado foi elaborado com base no Sistema Nacional de Pesquisa de Custos e Índices da Construção Civil (SINAPI), referência {{{{DATA_BASE}}}}, com aplicação de BDI de {{{{BDI}}}}%.

O cronograma físico-financeiro e as planilhas orçamentárias detalhadas encontram-se em anexo."""
    
    p.add_run(texto_orcamento)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Seção 6: Considerações Finais
    doc.add_heading("6. CONSIDERAÇÕES FINAIS", level=2)
    p = doc.add_paragraph()
    texto_final = """Este Memorial Descritivo, em conjunto com os projetos executivos, planilhas orçamentárias e demais documentos técnicos, compõe o conjunto de elementos necessários para a perfeita execução do objeto.

Quaisquer dúvidas ou necessidades de esclarecimentos deverão ser dirigidas ao responsável técnico do projeto."""
    
    p.add_run(texto_final)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nova Petrópolis/RS, {{{{DATA_ATUAL}}}}\n\n")
    p.add_run("_" * 50 + "\n")
    p.add_run("{{RESPONSAVEL}}\n")
    p.add_run("{{SETOR}}")
    
    doc.save(caminho_saida)
    print(f"✅ Template MD criado: {caminho_saida}")

def main():
    """Gera todos os templates."""
    
    # Caminho da pasta de templates
    pasta_templates = "/root/pmnova/engenharia-portal/shared_files/templates"
    
    print("=" * 80)
    print("🏗️  GERADOR DE TEMPLATES WORD - DOCUMENTOS TÉCNICOS")
    print("=" * 80)
    print()
    
    # Criar templates
    criar_template_etp(os.path.join(pasta_templates, "template_etp.docx"))
    criar_template_tr(os.path.join(pasta_templates, "template_tr.docx"))
    criar_template_md(os.path.join(pasta_templates, "template_md.docx"))
    
    print()
    print("=" * 80)
    print("✅ TODOS OS TEMPLATES FORAM CRIADOS COM SUCESSO!")
    print("=" * 80)
    print()
    print("📂 Localização:", pasta_templates)
    print()
    print("📋 Placeholders disponíveis:")
    print("   • {{OBJETO}} - Descrição do objeto")
    print("   • {{VALOR_GLOBAL}} - Valor total")
    print("   • {{VALOR_REPASSE}} - Valor do repasse")
    print("   • {{VALOR_CONTRAPARTIDA}} - Contrapartida municipal")
    print("   • {{AREA_TOTAL}} - Área em m²")
    print("   • {{BDI}} - BDI em %")
    print("   • {{DATA_BASE}} - Data base do orçamento")
    print("   • {{LOCAL}} - Município/UF")
    print("   • {{RESPONSAVEL}} - Nome do responsável")
    print("   • {{SETOR}} - Setor/Secretaria")
    print("   • {{TEXTO_IA}} - Texto gerado pela IA")
    print("   • {{JUSTIFICATIVA}} - Justificativa complementar")
    print("   • {{DATA_ATUAL}} - Data de geração")
    print()

if __name__ == "__main__":
    main()
