"""
Script de exemplo: Gerador de Peças Técnicas
Utilizado pelo n8n para processar e gerar documentos
"""

import os
import sys
from datetime import datetime
from docx import Document
from typing import Dict, Any
import json


def substituir_variaveis_docx(template_path: str, output_path: str, variaveis: Dict[str, Any]) -> str:
    """
    Substitui variáveis em um template .docx
    
    Args:
        template_path: Caminho do template
        output_path: Caminho do arquivo de saída
        variaveis: Dicionário com as variáveis a substituir
    
    Returns:
        Caminho do arquivo gerado
    """
    try:
        # Carregar o template
        doc = Document(template_path)
        
        # Substituir em parágrafos
        for paragrafo in doc.paragraphs:
            for chave, valor in variaveis.items():
                placeholder = f"{{{{{chave}}}}}"
                if placeholder in paragrafo.text:
                    # Substituir mantendo a formatação
                    for run in paragrafo.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(valor))
        
        # Substituir em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for chave, valor in variaveis.items():
                        placeholder = f"{{{{{chave}}}}}"
                        if placeholder in celula.text:
                            celula.text = celula.text.replace(placeholder, str(valor))
        
        # Salvar o documento
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        raise Exception(f"Erro ao processar documento: {str(e)}")


def gerar_peca_tecnica(tipo_peca: str, dados: Dict[str, Any]) -> str:
    """
    Gera uma peça técnica (ETP, TR ou MD) baseado nos dados fornecidos
    
    Args:
        tipo_peca: Tipo da peça (etp, tr, md)
        dados: Dados do formulário + texto gerado pela IA
    
    Returns:
        Caminho do arquivo gerado
    """
    # Definir caminhos
    template_dir = "/data/templates"
    output_dir = "/data/output"
    
    # Mapear tipo de peça para template
    templates = {
        "etp": "template_etp.docx",
        "tr": "template_tr.docx",
        "md": "template_md.docx"
    }
    
    if tipo_peca.lower() not in templates:
        raise ValueError(f"Tipo de peça inválido: {tipo_peca}")
    
    template_path = os.path.join(template_dir, templates[tipo_peca.lower()])
    
    # Verificar se template existe
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")
    
    # Preparar nome do arquivo de saída
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_arquivo = f"{tipo_peca.upper()}_{dados.get('objeto_resumido', 'documento')}_{timestamp}.docx"
    nome_arquivo = nome_arquivo.replace(" ", "_").replace("/", "-")
    output_path = os.path.join(output_dir, nome_arquivo)
    
    # Adicionar data e timestamp às variáveis
    dados_completos = {
        **dados,
        "DATA_ATUAL": datetime.now().strftime("%d/%m/%Y"),
        "HORA_ATUAL": datetime.now().strftime("%H:%M"),
        "ANO_ATUAL": datetime.now().strftime("%Y")
    }
    
    # Gerar documento
    arquivo_gerado = substituir_variaveis_docx(template_path, output_path, dados_completos)
    
    return arquivo_gerado


if __name__ == "__main__":
    # Exemplo de uso via linha de comando
    # python gerar_peca.py '{"tipo": "etp", "dados": {...}}'
    
    if len(sys.argv) > 1:
        try:
            parametros = json.loads(sys.argv[1])
            tipo = parametros.get("tipo", "etp")
            dados = parametros.get("dados", {})
            
            resultado = gerar_peca_tecnica(tipo, dados)
            print(json.dumps({"status": "success", "arquivo": resultado}))
            
        except Exception as e:
            print(json.dumps({"status": "error", "mensagem": str(e)}))
            sys.exit(1)
    else:
        print("Uso: python gerar_peca.py '{\"tipo\": \"etp\", \"dados\": {...}}'")
